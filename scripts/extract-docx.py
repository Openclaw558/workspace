#!/usr/bin/env python3
"""
Extract knowledgeChronicle.docx → structured markdown files
Output: workspace/docs/knowledge-base/chronicle/
"""

import docx
import os
import re
import json

DOCX_PATH = "/Users/ahmadfaris/.openclaw/docs/knowledgeChronicle.docx"
OUTPUT_DIR = "/Users/ahmadfaris/.openclaw/workspace/docs/knowledge-base/chronicle"

def slugify(text):
    """Convert text to a filename-friendly slug"""
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')[:80]

def extract_table(table):
    """Extract table data as markdown"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    # Build markdown table
    lines = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row if needed
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
    
    return "\n".join(lines)

def main():
    doc = docx.Document(DOCX_PATH)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Track sections
    current_role = None
    current_screen = None
    current_subsection = None  # Narration, Mermaid Chart, User Journey Table
    sections = []
    current_section = None
    
    # We'll parse paragraph by paragraph, tracking role + screen
    table_idx = 0
    para_to_table = {}  # map paragraph index to table
    
    # Build full content with tables interleaved
    # First, get all content blocks
    content_blocks = []
    
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            # Find matching paragraph
            for i, p in enumerate(doc.paragraphs):
                if p._element is element:
                    content_blocks.append(('paragraph', i, p))
                    break
        elif tag == 'tbl':
            # Find matching table
            for i, t in enumerate(doc.tables):
                if t._element is element:
                    content_blocks.append(('table', i, t))
                    break
    
    # Now process blocks sequentially
    all_sections = []
    current_role = "unknown"
    current_screen = "intro"
    section_content = []
    section_meta = {"role": current_role, "screen": current_screen}
    
    for block_type, idx, obj in content_blocks:
        if block_type == 'paragraph':
            p = obj
            text = p.text.strip()
            style = p.style.name if p.style else "Normal"
            
            if not text:
                continue
            
            # Detect role headers
            role_match = re.match(r'^USER JOURNEY ROLE\s+(\w+)', text, re.IGNORECASE)
            if role_match or style in ('Heading 1', 'Title') and 'USER JOURNEY' in text.upper():
                if role_match:
                    new_role = role_match.group(1).lower()
                elif 'MANAGER' in text.upper():
                    new_role = 'manager'
                elif 'ADMIN' in text.upper():
                    new_role = 'admin'
                elif 'OWNER' in text.upper():
                    new_role = 'owner'
                else:
                    new_role = slugify(text)
                
                # Save previous section
                if section_content:
                    all_sections.append({
                        "role": current_role,
                        "screen": current_screen,
                        "content": "\n".join(section_content)
                    })
                
                current_role = new_role
                current_screen = "intro"
                section_content = [f"# User Journey: Role {new_role.upper()}\n"]
                continue
            
            # Detect screen names (they're typically short, bold, and not narration/mermaid)
            is_screen_header = False
            if style in ('Heading 2', 'Heading 3'):
                is_screen_header = True
            elif (len(text) < 80 and 
                  text not in ('Narration:', 'Narrative:', 'Mermaid Chart:', 'User Journey Table:', 'User Journey Table', 'User Journey table') and
                  not text.startswith('The ') and 
                  not text.startswith('Upon ') and
                  not text.startswith('Beyond ') and
                  'narration' not in text.lower() and
                  'mermaid' not in text.lower()):
                # Check if runs are bold
                all_bold = all(run.bold for run in p.runs if run.text.strip()) if p.runs else False
                if all_bold and p.runs:
                    is_screen_header = True
                # Also detect screen names based on known patterns
                screen_patterns = [
                    'login', 'dashboard', 'map', 'tables', 'calendar', 'request', 
                    'sales', 'report', 'organization', 'profile', 'general',
                    'cemeteries', 'access', 'custom field', 'custom sales',
                    'event type', 'business type', 'regional', 'certificates', 'forms',
                    'invitation', 'onboarding', 'toolbar', 'sidebar', 'help', 'about',
                    'log out', 'logout'
                ]
                text_lower = text.lower()
                for pattern in screen_patterns:
                    if pattern in text_lower and len(text) < 100:
                        is_screen_header = True
                        break
            
            if is_screen_header and not any(kw in text.lower() for kw in ['narration', 'narrative', 'mermaid chart', 'user journey']):
                # Save previous section
                if section_content:
                    all_sections.append({
                        "role": current_role,
                        "screen": current_screen,
                        "content": "\n".join(section_content)
                    })
                
                current_screen = text
                section_content = [f"\n## {text}\n"]
                continue
            
            # Regular content - add to current section
            if 'Narration' in text or 'Narrative' in text:
                section_content.append(f"\n### Narration\n")
            elif 'Mermaid Chart' in text:
                section_content.append(f"\n### Mermaid Chart\n")
            elif 'User Journey Table' in text or 'User Journey table' in text:
                section_content.append(f"\n### User Journey Table\n")
            else:
                section_content.append(text)
        
        elif block_type == 'table':
            table = obj
            md_table = extract_table(table)
            if md_table:
                section_content.append(f"\n{md_table}\n")
    
    # Save last section
    if section_content:
        all_sections.append({
            "role": current_role,
            "screen": current_screen,
            "content": "\n".join(section_content)
        })
    
    # Write files per role
    role_files = {}
    for section in all_sections:
        role = section["role"]
        if role not in role_files:
            role_files[role] = []
        role_files[role].append(section)
    
    # Write individual role files
    index_entries = []
    for role, sections in role_files.items():
        filename = f"role-{role}.md"
        filepath = os.path.join(OUTPUT_DIR, filename)
        
        with open(filepath, 'w') as f:
            f.write(f"# Chronicle Knowledge Base — Role: {role.upper()}\n\n")
            f.write(f"_Extracted from knowledgeChronicle.docx_\n\n")
            f.write("---\n\n")
            
            screens = []
            for section in sections:
                f.write(section["content"])
                f.write("\n\n---\n\n")
                screens.append(section["screen"])
            
            index_entries.append({
                "role": role,
                "file": filename,
                "screens": screens
            })
        
        print(f"  ✅ {filepath} ({len(sections)} sections)")
    
    # Write a combined knowledge-index.json for retrieval
    index_path = os.path.join(OUTPUT_DIR, "knowledge-index.json")
    with open(index_path, 'w') as f:
        json.dump({
            "source": "knowledgeChronicle.docx",
            "extractedAt": "2026-02-20",
            "roles": index_entries,
            "totalSections": len(all_sections),
            "screens": list(set(s["screen"] for s in all_sections))
        }, f, indent=2)
    print(f"  ✅ {index_path}")
    
    # Also write a single combined file for full-text search
    combined_path = os.path.join(OUTPUT_DIR, "chronicle-full.md")
    with open(combined_path, 'w') as f:
        f.write("# Chronicle — Complete Knowledge Base\n\n")
        f.write("_Digital Cemetery Management Platform — All User Journeys_\n\n")
        f.write("---\n\n")
        for section in all_sections:
            f.write(section["content"])
            f.write("\n\n---\n\n")
    print(f"  ✅ {combined_path}")
    
    print(f"\n✅ Done! Extracted {len(all_sections)} sections for {len(role_files)} roles")

if __name__ == "__main__":
    main()
